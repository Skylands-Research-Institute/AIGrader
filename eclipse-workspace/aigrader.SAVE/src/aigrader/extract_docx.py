# extract_docx.py
"""
DOCX → plain-text extraction for AIGrader.

Design goals:
- Deterministic, boring, and testable.
- Extracts human-readable content in reading order (best-effort).
- Ignores layout/formatting (margins, fonts, spacing) by design.
- Works with either a filesystem path or raw bytes (downloaded from Canvas).

Notes / limitations:
- Text inside floating text boxes/shapes may not be captured reliably.
- Footnotes/endnotes/header/footer require extra handling; optional helpers included.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import Iterable, List, Optional, Sequence, Union

from docx import Document  # pip install python-docx


@dataclass(frozen=True)
class DocxExtractResult:
    text: str
    paragraphs_extracted: int
    table_cells_extracted: int
    has_header_text: bool
    has_footer_text: bool


def extract_docx_text(
    docx: Union[str, bytes, BytesIO],
    *,
    include_tables: bool = True,
    include_headers_footers: bool = False,
    keep_blank_lines: bool = False,
    table_cell_delim: str = " | ",
    table_row_delim: str = "\n",
) -> DocxExtractResult:
    """
    Extract plain text from a .docx file.

    Args:
        docx: Path to .docx OR raw bytes OR BytesIO.
        include_tables: If True, includes table contents in output.
        include_headers_footers: If True, include header/footer text (best-effort).
        keep_blank_lines: If True, preserve blank paragraphs as blank lines.
        table_cell_delim: Delimiter between cells when flattening a table row.
        table_row_delim: Delimiter between rows when flattening a table.

    Returns:
        DocxExtractResult with extracted text plus basic counters.

    Raises:
        ValueError if input is empty or not a docx.
    """
    document = _load_document(docx)

    parts: List[str] = []
    para_count = 0
    cell_count = 0

    # Paragraphs in body
    for p in document.paragraphs:
        para_count += 1
        t = (p.text or "").strip()
        if t:
            parts.append(t)
        elif keep_blank_lines:
            parts.append("")

    # Tables (optional)
    if include_tables:
        t_parts, extracted_cells = _extract_tables_text(
            document,
            cell_delim=table_cell_delim,
            row_delim=table_row_delim,
        )
        if t_parts:
            # Separate body text from tables a bit for readability
            if parts and parts[-1] != "":
                parts.append("")
            parts.extend(t_parts)
        cell_count += extracted_cells

    has_header = False
    has_footer = False

    # Header/footer (optional, best-effort)
    if include_headers_footers:
        header_text = _extract_headers_text(document)
        footer_text = _extract_footers_text(document)
        if header_text.strip():
            has_header = True
            parts = [header_text.strip(), ""] + parts
        if footer_text.strip():
            has_footer = True
            if parts and parts[-1] != "":
                parts.append("")
            parts.append(footer_text.strip())

    # Normalize whitespace: keep paragraph breaks but avoid huge runs of blanks
    text = _normalize_parts(parts)

    return DocxExtractResult(
        text=text,
        paragraphs_extracted=para_count,
        table_cells_extracted=cell_count,
        has_header_text=has_header,
        has_footer_text=has_footer,
    )


# -------------------------
# Internal helpers
# -------------------------

def _load_document(docx: Union[str, bytes, BytesIO]) -> Document:
    if isinstance(docx, str):
        # assume file path
        return Document(docx)
    if isinstance(docx, bytes):
        if not docx:
            raise ValueError("Empty DOCX bytes.")
        return Document(BytesIO(docx))
    if isinstance(docx, BytesIO):
        # Ensure position at start
        docx.seek(0)
        return Document(docx)
    raise ValueError(f"Unsupported docx input type: {type(docx)}")


def _extract_tables_text(
    document: Document,
    *,
    cell_delim: str,
    row_delim: str,
) -> tuple[List[str], int]:
    """
    Flatten tables into readable lines.
    Returns (lines, cell_count).
    """
    lines: List[str] = []
    cell_count = 0

    for table in document.tables:
        for row in table.rows:
            # Some docx files duplicate cell objects in merged cells;
            # still OK for plain text usage.
            cells = [(_cell_text(c) or "").strip() for c in row.cells]
            cell_count += len(row.cells)

            # Drop fully empty rows
            if all(not c for c in cells):
                continue

            # Keep non-empty cells; preserve column breaks
            compact = [c for c in cells if c]
            line = cell_delim.join(compact) if compact else ""
            if line:
                lines.append(line)

        # blank line between tables
        if lines and lines[-1] != "":
            lines.append("")

    # If we ended with an extra blank line, remove it
    while lines and lines[-1] == "":
        lines.pop()

    return lines, cell_count


def _cell_text(cell) -> str:
    # A cell can contain multiple paragraphs
    texts: List[str] = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            texts.append(t)
    return "\n".join(texts)


def _extract_headers_text(document: Document) -> str:
    """
    Extract header text from all sections (best-effort).
    """
    texts: List[str] = []
    for section in document.sections:
        header = section.header
        for p in header.paragraphs:
            t = (p.text or "").strip()
            if t:
                texts.append(t)
    return "\n".join(_dedupe_preserve_order(texts))


def _extract_footers_text(document: Document) -> str:
    """
    Extract footer text from all sections (best-effort).
    """
    texts: List[str] = []
    for section in document.sections:
        footer = section.footer
        for p in footer.paragraphs:
            t = (p.text or "").strip()
            if t:
                texts.append(t)
    return "\n".join(_dedupe_preserve_order(texts))


def _dedupe_preserve_order(items: Sequence[str]) -> List[str]:
    seen = set()
    out: List[str] = []
    for it in items:
        if it not in seen:
            out.append(it)
            seen.add(it)
    return out


def _normalize_parts(parts: Sequence[str]) -> str:
    """
    Join parts with blank lines between paragraphs, and collapse
    runs of >2 blank lines down to 2.
    """
    # Join with double newlines to represent paragraph breaks.
    raw = "\n\n".join(parts)

    # Collapse excessive blank lines while preserving paragraph breaks.
    lines = raw.splitlines()
    normalized: List[str] = []
    blank_run = 0

    for line in lines:
        if line.strip() == "":
            blank_run += 1
            # Allow at most 2 consecutive blank lines
            if blank_run <= 2:
                normalized.append("")
        else:
            blank_run = 0
            normalized.append(line.rstrip())

    # Strip leading/trailing blank lines
    while normalized and normalized[0] == "":
        normalized.pop(0)
    while normalized and normalized[-1] == "":
        normalized.pop()

    return "\n".join(normalized)
